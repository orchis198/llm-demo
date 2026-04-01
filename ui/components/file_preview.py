from __future__ import annotations

from pathlib import Path
from typing import Any

import pandas as pd
import streamlit as st


def _extract_docx_text(file_path: str) -> str:
    from docx import Document

    document = Document(file_path)
    lines: list[str] = []
    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if text:
            lines.append(text)
    for table in document.tables:
        for row in table.rows:
            row_text = " ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
            if row_text:
                lines.append(row_text)
    return "\n".join(lines).strip() or "当前文档暂无可提取文本。"


def _render_file_preview_content(file_path: str, *, height: int, key_prefix: str) -> None:
    suffix = Path(file_path).suffix.lower()
    file_name = Path(file_path).name
    if suffix in {".txt", ".md"}:
        st.text_area("文件内容", Path(file_path).read_text(encoding="utf-8"), height=height, disabled=True, key=f"{key_prefix}_{file_name}_text")
        return
    if suffix in {".png", ".jpg", ".jpeg", ".avif"}:
        st.image(file_path, use_container_width=True)
        return
    if suffix in {".xls", ".xlsx"}:
        workbook = pd.read_excel(file_path, sheet_name=None, dtype=str, engine=None)
        if not workbook:
            st.info("当前表格没有可展示内容。")
            return
        sheet_names = list(workbook.keys())
        selected_sheet = st.selectbox("工作表", sheet_names, key=f"{key_prefix}_{file_name}_sheet")
        st.dataframe(workbook[selected_sheet].fillna(""), use_container_width=True, hide_index=True, height=height)
        return
    if suffix == ".docx":
        try:
            content = _extract_docx_text(file_path)
            st.text_area("文档内容", content, height=height, disabled=True, key=f"{key_prefix}_{file_name}_docx")
        except Exception as exc:
            st.info(f"当前文档暂不支持完整预览：{exc}")
        return
    st.info("当前文件类型暂不支持预览。")


def render_file_preview(file_path: str, *, height: int = 420, allow_expand: bool = False) -> None:
    _render_file_preview_content(file_path, height=height, key_prefix="preview")
    if allow_expand:
        with st.popover("放大预览"):
            _render_file_preview_content(file_path, height=max(height + 260, 760), key_prefix="preview_expand")


def render_source_snapshot_preview(
    container,
    parsed_invoice: dict[str, Any],
    *,
    height: int = 480,
) -> None:
    """在给定 container（可以是 st.sidebar 或 st.container）内渲染原件快照预览。

    支持两种模式：
    - PDF：渲染分页图（需调用方预先把页图存入 st.session_state.source_page_images）
    - 文本/其他：渲染 text_area 显示原始文本
    """
    source_file = parsed_invoice.get("source_file", "") or "未记录"
    container.caption(f"来源文件：{source_file}")

    source_page_images: list[bytes] | list = st.session_state.get("source_page_images") or []

    if parsed_invoice.get("source_type") == "pdf" and source_page_images:
        total_pages = len(source_page_images)
        current_index = min(max(int(st.session_state.get("preview_page_index", 0)), 0), total_pages - 1)

        nav_cols = container.columns([1, 1, 1.3])
        if nav_cols[0].button(
            "上一页",
            disabled=current_index <= 0,
            key="snap_prev_page",
            use_container_width=True,
        ):
            st.session_state.preview_page_index = current_index - 1
            st.rerun()
        if nav_cols[1].button(
            "下一页",
            disabled=current_index >= total_pages - 1,
            key="snap_next_page",
            use_container_width=True,
        ):
            st.session_state.preview_page_index = current_index + 1
            st.rerun()
        selected_page = nav_cols[2].number_input(
            "页码",
            min_value=1,
            max_value=total_pages,
            value=current_index + 1,
            step=1,
            key="snap_page_number",
            label_visibility="collapsed",
        )
        if selected_page - 1 != current_index:
            st.session_state.preview_page_index = selected_page - 1
            st.rerun()
        container.caption(f"第 {current_index + 1} / {total_pages} 页")
        container.image(source_page_images[current_index], use_container_width=True)
        return

    raw_text = parsed_invoice.get("raw_text", "") or "(无原始文本)"
    container.text_area("原始文本", raw_text, height=height, disabled=True, label_visibility="collapsed")
