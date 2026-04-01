from __future__ import annotations

import copy
import json
import os
import tempfile
from io import BytesIO
from pathlib import Path
from typing import Any

import pandas as pd

from engines.invoice_parser import InvoiceParser

SUPPORTED_SUFFIXES = (".pdf", ".txt", ".docx", ".xlsx", ".xls")

HEADER_FIELDS = [
    ("invoice_title", "票据标题"),
    ("invoice_type", "发票类型"),
    ("invoice_number", "发票号码"),
    ("issue_date", "开票日期"),
    ("check_code", "校验码"),
    ("machine_number", "机器编号"),
    ("buyer_name", "购方名称"),
    ("buyer_tax_id", "购方纳税人识别号"),
    ("buyer_address_phone", "购方地址、电话"),
    ("buyer_bank_account", "购方开户行及账号"),
    ("seller_name", "销方名称"),
    ("seller_tax_id", "销方纳税人识别号"),
    ("seller_address_phone", "销方地址、电话"),
    ("seller_bank_account", "销方开户行及账号"),
    ("subtotal_amount", "金额合计"),
    ("tax_amount", "税额合计"),
    ("total_amount", "价税合计"),
    ("total_amount_chinese", "价税合计大写"),
    ("tax_classification_code", "税收分类编码"),
    ("remarks", "备注"),
    ("payee", "收款人"),
    ("reviewer", "复核"),
    ("issuer", "开票人"),
    ("invoice_status", "发票状态"),
    ("original_invoice_code", "原发票代码"),
    ("original_invoice_number", "原发票号码"),
]

NUMERIC_HEADER_FIELDS = {"subtotal_amount", "tax_amount", "total_amount"}
FIELD_STATUS_OPTIONS = ["正常", "存疑", "有误", "待补充"]
CRITICAL_WARNING_STATUSES = {"存疑", "有误", "待补充"}
KEY_REVIEW_FIELDS = [
    "invoice_number",
    "issue_date",
    "buyer_name",
    "seller_name",
    "subtotal_amount",
    "tax_amount",
    "total_amount",
]
LONG_TEXT_FIELDS = {"remarks"}

LINE_ITEM_COLUMNS = [
    ("item_name", "项目名称"),
    ("spec_model", "规格型号"),
    ("unit", "单位"),
    ("quantity", "数量"),
    ("unit_price", "单价"),
    ("amount", "金额"),
    ("tax_rate", "税率"),
    ("tax_amount", "税额"),
]

LINE_ITEM_NUMERIC_FIELDS = {"quantity", "unit_price", "amount", "tax_amount"}
LINE_ITEM_EDITOR_COLUMNS = [column_label for _, column_label in LINE_ITEM_COLUMNS] + ["状态"]

REVIEW_OWNER_FIELD = "review_owner"


# ---------------------------------------------------------------------------
# 示例文件收集
# ---------------------------------------------------------------------------


def collect_sample_files(base_dir: str) -> dict[str, str]:
    base = Path(base_dir).resolve()
    sample_files: dict[str, str] = {}
    for folder_name in ("data", "example"):
        folder_path = base / folder_name
        if not folder_path.exists():
            continue
        for file_path in folder_path.rglob("*"):
            if not file_path.is_file() or file_path.suffix.lower() not in SUPPORTED_SUFFIXES:
                continue
            rel = file_path.relative_to(base).as_posix()
            sample_files[rel] = str(file_path)
    return dict(sorted(sample_files.items()))


# ---------------------------------------------------------------------------
# 文件内容读取
# ---------------------------------------------------------------------------


def decode_text_bytes(file_bytes: bytes) -> str:
    for enc in ("utf-8-sig", "utf-8", "gb18030"):
        try:
            return file_bytes.decode(enc)
        except UnicodeDecodeError:
            continue
    return file_bytes.decode("utf-8", errors="ignore")


def extract_docx_text_from_bytes(file_bytes: bytes) -> str:
    from docx import Document

    document = Document(BytesIO(file_bytes))
    lines: list[str] = []
    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if text:
            lines.append(text)
    for table in document.tables:
        for row in table.rows:
            normalized = " ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
            if normalized:
                lines.append(normalized)
    return "\n".join(lines).strip()


def extract_excel_text_from_bytes(file_bytes: bytes) -> str:
    sheets = pd.read_excel(BytesIO(file_bytes), sheet_name=None, header=None, dtype=str, engine=None)
    return _extract_excel_text(sheets)


def _extract_excel_text(sheets: dict) -> str:
    lines: list[str] = []
    for sheet_name, df in sheets.items():
        for _, row in df.iterrows():
            normalized = " ".join(str(v).strip() for v in row.values if str(v).strip())
            if normalized:
                lines.append(normalized)
    return "\n".join(lines).strip()


def render_pdf_pages(file_bytes: bytes) -> list[bytes]:
    import io

    import pdfplumber

    page_images: list[bytes] = []
    with pdfplumber.open(BytesIO(file_bytes)) as pdf:
        for page in pdf.pages:
            image = page.to_image(resolution=140).original
            buffer = io.BytesIO()
            image.save(buffer, format="PNG")
            page_images.append(buffer.getvalue())
    return page_images


# ---------------------------------------------------------------------------
# 发票解析
# ---------------------------------------------------------------------------


def parse_invoice(
    display_name: str,
    file_path: str | None = None,
    file_bytes: bytes | None = None,
) -> dict[str, Any]:
    parser = InvoiceParser()
    suffix = Path(display_name).suffix.lower()

    if suffix == ".txt":
        if file_bytes is not None:
            text = decode_text_bytes(file_bytes)
        else:
            text = decode_text_bytes(Path(file_path or "").read_bytes())
        invoice_data = parser.parse_text(text)
        invoice_data["source_type"] = "txt"
        invoice_data["source_file"] = display_name
        return invoice_data

    if suffix == ".pdf":
        if file_path:
            invoice_data = parser.parse_pdf(file_path)
        else:
            tmp = _write_temp_file(file_bytes or b"", ".pdf")
            try:
                invoice_data = parser.parse_pdf(tmp)
            finally:
                if os.path.exists(tmp):
                    os.unlink(tmp)
        invoice_data["source_type"] = "pdf"
        invoice_data["source_file"] = display_name
        return invoice_data

    if suffix == ".docx":
        if file_bytes is not None:
            text = extract_docx_text_from_bytes(file_bytes)
        else:
            from docx import Document

            doc = Document(file_path or "")
            lines: list[str] = []
            for para in doc.paragraphs:
                if para.text.strip():
                    lines.append(para.text.strip())
            text = "\n".join(lines)
        invoice_data = parser.parse_text(text)
        invoice_data["source_type"] = "docx"
        invoice_data["source_file"] = display_name
        return invoice_data

    # .xlsx / .xls
    if file_bytes is not None:
        text = extract_excel_text_from_bytes(file_bytes)
    else:
        sheets = pd.read_excel(file_path or "", sheet_name=None, header=None, dtype=str, engine=None)
        text = _extract_excel_text(sheets)
    invoice_data = parser.parse_text(text)
    invoice_data["source_type"] = suffix.lstrip(".")
    invoice_data["source_file"] = display_name
    return invoice_data


def _write_temp_file(file_bytes: bytes, suffix: str) -> str:
    with tempfile.NamedTemporaryFile(delete=False, suffix=suffix) as f:
        f.write(file_bytes)
        return f.name


# ---------------------------------------------------------------------------
# 解析完成后填充 session state（由 page 层调用）
# ---------------------------------------------------------------------------


def apply_parsed_invoice(
    parsed_invoice: dict[str, Any],
    source_file_bytes: bytes | None,
    source_page_images: list[bytes],
) -> None:
    import streamlit as st

    st.session_state.parsed_invoice = parsed_invoice
    st.session_state.source_file_bytes = source_file_bytes
    st.session_state.source_page_images = source_page_images
    st.session_state.preview_page_index = 0
    st.session_state.review_draft = build_review_draft(parsed_invoice)
    st.session_state.saved_review = None
    st.session_state.submitted_review = None
    st.session_state.verify_result = None


# ---------------------------------------------------------------------------
# 审核草稿构建与定稿
# ---------------------------------------------------------------------------


def build_review_draft(invoice_data: dict[str, Any]) -> dict[str, Any]:
    draft: dict[str, Any] = {
        "field_values": {},
        "field_statuses": {},
        "line_items": [],
        REVIEW_OWNER_FIELD: "",
    }
    normalized = copy.deepcopy(invoice_data)
    if not normalized.get("invoice_number") and normalized.get("invoice_code"):
        normalized["invoice_number"] = normalized.get("invoice_code")
    for field_key, _ in HEADER_FIELDS:
        draft["field_values"][field_key] = normalized.get(field_key)
        draft["field_statuses"][field_key] = "正常" if normalized.get(field_key) not in (None, "", []) else "待补充"
    for record in invoice_data.get("line_items", []):
        draft["line_items"].append({
            "values": {column_key: record.get(column_key, "") or "" for column_key, _ in LINE_ITEM_COLUMNS},
            "status": "正常",
        })
    if not draft["line_items"]:
        draft["line_items"].append({"values": {column_key: "" for column_key, _ in LINE_ITEM_COLUMNS}, "status": "待补充"})
    return draft


def normalize_line_item_record(record: dict[str, Any]) -> dict[str, Any]:
    return {column_key: (record.get(column_key, "") or "") for column_key, _ in LINE_ITEM_COLUMNS}


def _to_display_text(value: Any) -> str:
    if value is None:
        return ""
    if isinstance(value, float):
        return f"{value:.2f}".rstrip("0").rstrip(".") if not value.is_integer() else str(int(value))
    if isinstance(value, list):
        return json.dumps(value, ensure_ascii=False)
    return str(value)


def _coerce_number(value: Any) -> float | None:
    if value in (None, ""):
        return None
    if isinstance(value, float) and pd.isna(value):
        return None
    if isinstance(value, (int, float)):
        return float(value)
    text = str(value).replace(",", "").strip()
    if not text:
        return None
    try:
        return float(text)
    except ValueError:
        return None


def _sanitize_text(value: Any) -> str:
    if value is None:
        return ""
    if isinstance(value, float) and pd.isna(value):
        return ""
    return str(value).strip()


def finalize_review_draft(review_draft: dict[str, Any], parsed_invoice: dict[str, Any]) -> dict[str, Any]:
    finalized = copy.deepcopy(parsed_invoice)
    finalized_statuses: dict[str, str] = {}
    finalized[REVIEW_OWNER_FIELD] = _sanitize_text(review_draft.get(REVIEW_OWNER_FIELD, ""))
    for field_key, _ in HEADER_FIELDS:
        raw_value = review_draft["field_values"].get(field_key)
        if field_key in NUMERIC_HEADER_FIELDS:
            finalized[field_key] = _coerce_number(raw_value)
        else:
            finalized[field_key] = _sanitize_text(raw_value)
        finalized_statuses[field_key] = review_draft["field_statuses"].get(field_key, "正常")

    finalized_line_items: list[dict[str, Any]] = []
    line_item_statuses: list[str] = []
    for record in review_draft.get("line_items", []):
        values = record.get("values", {})
        normalized_record: dict[str, Any] = {}
        has_content = False
        for column_key, _ in LINE_ITEM_COLUMNS:
            raw_value = values.get(column_key, "")
            if column_key in LINE_ITEM_NUMERIC_FIELDS:
                normalized_value = _coerce_number(raw_value)
            else:
                normalized_value = _sanitize_text(raw_value)
            if normalized_value not in (None, ""):
                has_content = True
            normalized_record[column_key] = normalized_value
        if has_content or any(str(v).strip() for v in normalized_record.values()):
            finalized_line_items.append(normalized_record)
            line_item_statuses.append(record.get("status", "正常"))

    finalized["line_items"] = finalized_line_items
    finalized["line_item_count"] = len(finalized_line_items)
    finalized["tax_rates"] = sorted(
        {str(item.get("tax_rate", "")).strip() for item in finalized_line_items if str(item.get("tax_rate", "")).strip()}
    )
    finalized["review_field_statuses"] = finalized_statuses
    finalized["review_line_item_statuses"] = line_item_statuses
    return finalized


# ---------------------------------------------------------------------------
# 差异对比
# ---------------------------------------------------------------------------


def collect_field_changes(parsed_invoice: dict[str, Any], finalized_review: dict[str, Any]) -> list[dict[str, str]]:
    changes: list[dict[str, str]] = []
    review_owner = _to_display_text(finalized_review.get(REVIEW_OWNER_FIELD))
    if review_owner:
        changes.append({"字段": "审核人", "原解析值": "", "审核值": review_owner, "状态": "正常"})
    for field_key, field_label in HEADER_FIELDS:
        before_value = _to_display_text(parsed_invoice.get(field_key))
        after_value = _to_display_text(finalized_review.get(field_key))
        status = _to_display_text(finalized_review.get("review_field_statuses", {}).get(field_key, "正常"))
        if before_value != after_value or status != "正常":
            changes.append({"字段": field_label, "原解析值": before_value, "审核值": after_value, "状态": status})

    original_line_items = json.dumps(parsed_invoice.get("line_items", []), ensure_ascii=False)
    reviewed_line_items = json.dumps(finalized_review.get("line_items", []), ensure_ascii=False)
    reviewed_statuses = finalized_review.get("review_line_item_statuses", [])
    if original_line_items != reviewed_line_items or reviewed_statuses:
        changes.append({
            "字段": "发票明细",
            "原解析值": original_line_items,
            "审核值": reviewed_line_items,
            "状态": json.dumps(reviewed_statuses, ensure_ascii=False),
        })
    return changes


def collect_submission_findings(finalized_review: dict[str, Any]) -> dict[str, Any]:
    field_statuses = finalized_review.get("review_field_statuses", {})
    critical_field_issues = [
        {"field_key": field_key, "label": field_label, "status": field_statuses.get(field_key, "正常")}
        for field_key, field_label in HEADER_FIELDS
        if field_key in KEY_REVIEW_FIELDS and field_statuses.get(field_key) in CRITICAL_WARNING_STATUSES
    ]
    optional_field_issues = [
        {"field_key": field_key, "label": field_label, "status": field_statuses.get(field_key, "正常")}
        for field_key, field_label in HEADER_FIELDS
        if field_key not in KEY_REVIEW_FIELDS and field_statuses.get(field_key) in CRITICAL_WARNING_STATUSES
    ]

    critical_line_item_issues: list[dict[str, Any]] = []
    line_items = finalized_review.get("line_items", [])
    for index, status in enumerate(finalized_review.get("review_line_item_statuses", []), start=1):
        if status not in CRITICAL_WARNING_STATUSES:
            continue
        line_item = line_items[index - 1] if index - 1 < len(line_items) else {}
        item_name = _sanitize_text(line_item.get("item_name")) or "未填写商品"
        critical_line_item_issues.append({"index": index, "item_name": item_name, "status": status})

    return {
        "critical_field_issues": critical_field_issues,
        "optional_field_issues": optional_field_issues,
        "critical_line_item_issues": critical_line_item_issues,
    }


def determine_review_conclusion(finalized_review: dict[str, Any]) -> str:
    findings = collect_submission_findings(finalized_review)
    critical_statuses = [issue["status"] for issue in findings["critical_field_issues"]] + [
        issue["status"] for issue in findings["critical_line_item_issues"]
    ]
    if any(status == "有误" for status in critical_statuses):
        return "审核不通过"
    if any(status in {"存疑", "待补充"} for status in critical_statuses):
        return "需补充确认"
    return "审核通过"


def validate_submission(finalized_review: dict[str, Any]) -> dict[str, Any]:
    errors: list[str] = []
    critical_messages: list[str] = []
    optional_messages: list[str] = []

    review_owner = _sanitize_text(finalized_review.get(REVIEW_OWNER_FIELD, ""))
    if not review_owner:
        errors.append("提交前必须填写审核人。")

    findings = collect_submission_findings(finalized_review)
    field_statuses = finalized_review.get("review_field_statuses", {})
    line_item_statuses = finalized_review.get("review_line_item_statuses", [])
    has_any_error = any(status == "有误" for status in field_statuses.values()) or any(status == "有误" for status in line_item_statuses)
    critical_statuses = [issue["status"] for issue in findings["critical_field_issues"]] + [
        issue["status"] for issue in findings["critical_line_item_issues"]
    ]

    if findings["critical_field_issues"]:
        critical_messages.append(
            "硬性字段仍存在异常："
            + "、".join(f"{issue['label']}（{issue['status']}）" for issue in findings["critical_field_issues"])
            + "。"
        )
    if findings["critical_line_item_issues"]:
        critical_messages.append(
            "关键明细仍存在异常："
            + "、".join(
                f"第{issue['index']}行 {issue['item_name']}（{issue['status']}）" for issue in findings["critical_line_item_issues"]
            )
            + "。"
        )
    if findings["optional_field_issues"]:
        optional_messages.append(
            "以下非硬性字段未直接阻断流程，但仍建议人工关注："
            + "、".join(f"{issue['label']}（{issue['status']}）" for issue in findings["optional_field_issues"])
            + "。"
        )
    if has_any_error:
        errors.append("当前存在字段或明细被标记为“有误”，已直接拦截，不能提交通过或确认采数结果。")

    conclusion = determine_review_conclusion(finalized_review)
    if has_any_error:
        conclusion = "审核不通过"
    if any(status in {"存疑", "待补充"} for status in critical_statuses) and not has_any_error:
        conclusion = "需补充确认"

    can_submit = not errors
    can_confirm = not errors and conclusion == "审核通过"
    return {
        "errors": errors,
        "critical_messages": critical_messages,
        "optional_messages": optional_messages,
        "conclusion": conclusion,
        "can_submit": can_submit,
        "can_confirm": can_confirm,
    }


# ---------------------------------------------------------------------------
# 示例数据判断
# ---------------------------------------------------------------------------


def is_sample_source(invoice_data: dict[str, Any]) -> bool:
    source_file = str(invoice_data.get("source_file", ""))
    return source_file.startswith("data/") or source_file.startswith("example/")


# ---------------------------------------------------------------------------
# 向后兼容：保留原有的 show 数据解析
# ---------------------------------------------------------------------------


def parse_show_invoice(text: str) -> dict[str, Any]:
    invoice_data = parse_invoice("data/show/raw/show_invoice.txt", file_bytes=text.encode("utf-8"))
    return invoice_data
